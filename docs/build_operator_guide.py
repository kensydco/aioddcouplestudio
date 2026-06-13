from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "AI_ODD_COUPLE_OPERATOR_GUIDE.md"
OUTPUT = ROOT / "docs" / "AI_Odd_Couple_Operator_Guide.docx"
BLUE = RGBColor(0x00, 0xA9, 0xFF)
RED = RGBColor(0xE6, 0x39, 0x46)
INK = RGBColor(0x1C, 0x1C, 0x20)
MUTED = RGBColor(0x60, 0x60, 0x68)
LIGHT = "F2F4F7"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_run(run, size=10.5, bold=False, color=INK, font="Arial", italic=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18
for name, size, color, before, after in [
    ("Heading 1", 16, RED, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, INK, 9, 4),
]:
    style = styles[name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("THE AI ODD COUPLE  |  OPERATOR GUIDE"), 8.5, True, MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(footer.add_run("Hosted production, approval, and social publishing"), 8, False, MUTED)

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(20)
cover.add_run().add_picture(str(ROOT / "assets" / "branding" / "AI Odd Couple Cover Page.png"), width=Inches(5.8))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
set_run(p.add_run("PRODUCTION SYSTEM"), 12, True, BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
set_run(p.add_run("Operator Guide"), 28, True, INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Capabilities, schedules, approvals, social publishing, and ad hoc prompts"), 11.5, False, MUTED, italic=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(26)
set_run(p.add_run("Version 1.0  |  June 13, 2026"), 9.5, True, RED)
doc.add_page_break()

lines = SOURCE.read_text(encoding="utf-8").splitlines()
i = 0
while i < len(lines):
    line = lines[i].strip()
    if not line or line.startswith("**Version:") or line.startswith("**Updated:") or line.startswith("**Repository:"):
        i += 1
        continue
    if line.startswith("# "):
        i += 1
        continue
    if line.startswith("## "):
        doc.add_heading(line[3:], level=1)
        i += 1
        continue
    if line.startswith("### "):
        doc.add_heading(line[4:], level=2)
        i += 1
        continue
    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        set_run(p.add_run(line[2:]), 10, False, INK, italic=True)
        i += 1
        continue
    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(line[2:]), 10.2)
        i += 1
        continue
    if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace("-", "").replace(":", "").strip()) == set():
        headers = [x.strip() for x in line.strip("|").split("|")]
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append([x.strip() for x in lines[i].strip().strip("|").split("|")])
            i += 1
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        set_repeat_table_header(table.rows[0])
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            shade(cell, "1C1C20")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_run(cell.paragraphs[0].add_run(text), 9.2, True, RGBColor(0xFF, 0xFF, 0xFF))
        for row_index, values in enumerate(rows):
            cells = table.add_row().cells
            for idx, text in enumerate(values):
                if row_index % 2 == 0:
                    shade(cells[idx], LIGHT)
                set_cell_margins(cells[idx])
                cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_run(cells[idx].paragraphs[0].add_run(text.replace("`", "")), 9)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        continue
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    text = line.replace("`", "")
    if text.startswith("1. **") or text.startswith("2. **"):
        set_run(p.add_run(text), 10.5, True)
    else:
        set_run(p.add_run(text), 10.5)
    i += 1

doc.core_properties.title = "The AI Odd Couple Production System Operator Guide"
doc.core_properties.subject = "Hosted production, approval, email notification, and social publishing"
doc.core_properties.author = "AI Odd Couple Studio"
doc.save(OUTPUT)
print(OUTPUT)
